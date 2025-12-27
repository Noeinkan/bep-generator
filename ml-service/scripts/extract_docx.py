"""
DOCX Text Extraction Script

Extracts text from DOCX files and saves them as plain text files
for the RAG system to process.
"""

import os
from pathlib import Path
from docx import Document
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_text_from_docx(docx_path):
    """
    Extract all text from a DOCX file

    Args:
        docx_path: Path to DOCX file

    Returns:
        Extracted text as string
    """
    try:
        doc = Document(docx_path)

        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(' | '.join(row_text))

        # Combine all text
        all_text = '\n'.join(paragraphs)
        if table_text:
            all_text += '\n\n' + '\n'.join(table_text)

        return all_text

    except Exception as e:
        logger.error(f"Error extracting from {docx_path}: {e}")
        return ""


def extract_all_docx(input_dir, output_dir):
    """
    Extract text from all DOCX files in a directory

    Args:
        input_dir: Directory containing DOCX files
        output_dir: Directory to save extracted text files
    """
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    docx_files = list(input_path.glob('*.docx'))
    logger.info(f"Found {len(docx_files)} DOCX files to process")

    for docx_file in docx_files:
        # Skip temporary Word files
        if docx_file.name.startswith('~$'):
            continue

        logger.info(f"Processing: {docx_file.name}")

        # Extract text
        text = extract_text_from_docx(docx_file)

        if text:
            # Save as .txt file
            txt_filename = docx_file.stem + '.txt'
            txt_path = output_path / txt_filename

            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(text)

            logger.info(f"  Saved: {txt_filename} ({len(text)} chars)")
        else:
            logger.warning(f"  No text extracted from {docx_file.name}")

    logger.info(f"Extraction complete! Files saved to: {output_path}")


def create_consolidated_file(txt_dir, output_file):
    """
    Consolidate all text files into one large training file

    Args:
        txt_dir: Directory containing extracted text files
        output_file: Path to consolidated output file
    """
    txt_path = Path(txt_dir)
    txt_files = list(txt_path.glob('*.txt'))

    logger.info(f"Consolidating {len(txt_files)} text files...")

    consolidated_text = []

    for txt_file in txt_files:
        with open(txt_file, 'r', encoding='utf-8') as f:
            content = f.read()
            # Add separator between documents
            consolidated_text.append(f"\n\n{'='*80}\n")
            consolidated_text.append(f"Document: {txt_file.name}\n")
            consolidated_text.append(f"{'='*80}\n\n")
            consolidated_text.append(content)

    # Write consolidated file
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(''.join(consolidated_text))

    total_chars = sum(len(t) for t in consolidated_text)
    logger.info(f"Consolidated file created: {output_path} ({total_chars} chars)")


if __name__ == "__main__":
    # Define paths
    script_dir = Path(__file__).parent
    ml_service_dir = script_dir.parent
    data_dir = ml_service_dir / 'data'

    docx_dir = data_dir / 'training_documents' / 'docx'
    txt_dir = data_dir / 'training_documents' / 'txt'
    consolidated_file = data_dir / 'consolidated_training_data.txt'

    logger.info("="*80)
    logger.info("DOCX Text Extraction for RAG System")
    logger.info("="*80)

    # Step 1: Extract individual DOCX files
    if docx_dir.exists():
        extract_all_docx(docx_dir, txt_dir)
    else:
        logger.error(f"DOCX directory not found: {docx_dir}")
        exit(1)

    # Step 2: Create consolidated file
    if txt_dir.exists():
        create_consolidated_file(txt_dir, consolidated_file)
    else:
        logger.warning(f"Text directory not found: {txt_dir}")

    logger.info("="*80)
    logger.info("Extraction Complete!")
    logger.info("="*80)
