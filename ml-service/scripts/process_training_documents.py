"""
Document Processor for BEP Training Data

This script extracts text from various document formats (PDF, DOCX, TXT)
and consolidates them into a single training corpus for the ML model.

Usage:
    python scripts/process_training_documents.py [--output OUTPUT_FILE]
"""

import os
import sys
import json
from pathlib import Path
from typing import List, Dict, Set
import argparse

# Document processing imports
import PyPDF2
import pdfplumber
from docx import Document
from tqdm import tqdm


class DocumentProcessor:
    """Process multiple document formats for ML training"""

    def __init__(self, data_dir: str = "data/training_documents"):
        # Make paths absolute relative to script location
        script_dir = Path(__file__).parent
        project_root = script_dir.parent

        self.data_dir = Path(data_dir)
        if not self.data_dir.is_absolute():
            self.data_dir = project_root / self.data_dir

        self.pdf_dir = self.data_dir / "pdf"
        self.docx_dir = self.data_dir / "docx"
        self.txt_dir = self.data_dir / "txt"

        self.project_root = project_root

        # Load glossary for context enhancement
        self.glossary = self._load_glossary()

        # Statistics
        self.stats = {
            'total_documents': 0,
            'pdf_count': 0,
            'docx_count': 0,
            'txt_count': 0,
            'total_chars': 0,
            'errors': []
        }

    def _load_glossary(self) -> Dict:
        """Load ISO 19650 glossary"""
        glossary_path = self.project_root / "data" / "iso19650_glossary.json"
        if glossary_path.exists():
            with open(glossary_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}

    def extract_text_from_pdf_pypdf2(self, pdf_path: Path) -> str:
        """Extract text from PDF using PyPDF2 (fallback method)"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            self.stats['errors'].append(f"PyPDF2 error in {pdf_path.name}: {str(e)}")
        return text

    def extract_text_from_pdf_pdfplumber(self, pdf_path: Path) -> str:
        """Extract text from PDF using pdfplumber (primary method)"""
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                    # Also extract text from tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
        except Exception as e:
            self.stats['errors'].append(f"pdfplumber error in {pdf_path.name}: {str(e)}")
            # Fallback to PyPDF2
            text = self.extract_text_from_pdf_pypdf2(pdf_path)
        return text

    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(docx_path)

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"

            # Extract headers/footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
        except Exception as e:
            self.stats['errors'].append(f"DOCX error in {docx_path.name}: {str(e)}")
        return text

    def extract_text_from_txt(self, txt_path: Path) -> str:
        """Extract text from TXT file"""
        text = ""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        except Exception as e:
            self.stats['errors'].append(f"TXT error in {txt_path.name}: {str(e)}")
        return text

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]

        # Remove duplicate consecutive lines
        cleaned_lines = []
        prev_line = None
        for line in lines:
            if line != prev_line:
                cleaned_lines.append(line)
            prev_line = line

        # Join with single newlines
        cleaned_text = "\n".join(cleaned_lines)

        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)

        return cleaned_text

    def process_directory(self, directory: Path, file_extension: str,
                         extraction_method) -> List[str]:
        """Process all files in a directory with given extension"""
        if not directory.exists():
            print(f"Directory {directory} does not exist, skipping...")
            return []

        files = list(directory.glob(f"*{file_extension}"))
        if not files:
            print(f"No {file_extension} files found in {directory}")
            return []

        extracted_texts = []
        print(f"\nProcessing {len(files)} {file_extension} files from {directory.name}...")

        for file_path in tqdm(files, desc=f"Processing {file_extension}"):
            try:
                text = extraction_method(file_path)
                cleaned_text = self.clean_text(text)

                if cleaned_text:
                    # Add document header
                    document_header = f"\n{'='*80}\n"
                    document_header += f"SOURCE: {file_path.name}\n"
                    document_header += f"{'='*80}\n\n"

                    extracted_texts.append(document_header + cleaned_text)
                    self.stats['total_chars'] += len(cleaned_text)
                    self.stats['total_documents'] += 1
                else:
                    self.stats['errors'].append(f"No text extracted from {file_path.name}")
            except Exception as e:
                self.stats['errors'].append(f"Failed to process {file_path.name}: {str(e)}")

        return extracted_texts

    def add_glossary_context(self) -> str:
        """Generate glossary-based training context"""
        if not self.glossary or 'terms' not in self.glossary:
            return ""

        context = "\n" + "="*80 + "\n"
        context += "ISO 19650 TERMINOLOGY REFERENCE\n"
        context += "="*80 + "\n\n"

        terms = self.glossary.get('terms', {})
        for term_key, term_data in terms.items():
            if 'usage_examples' in term_data:
                for example in term_data['usage_examples']:
                    context += example + "\n"

        # Add common phrases
        if 'common_phrases' in self.glossary:
            context += "\n" + "-"*80 + "\n"
            context += "Common BEP Phrases:\n"
            context += "-"*80 + "\n\n"

            for category, phrases in self.glossary['common_phrases'].items():
                for phrase in phrases[:3]:  # Limit to 3 examples per category
                    context += phrase + "\n"

        return context

    def process_all_documents(self, output_file: str = "data/consolidated_training_data.txt") -> str:
        """Process all documents and create consolidated training file"""
        print("="*80)
        print("BEP Training Document Processor")
        print("="*80)

        all_texts = []

        # Make output path absolute if needed
        output_path = Path(output_file)
        if not output_path.is_absolute():
            output_path = self.project_root / output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Process PDFs
        pdf_texts = self.process_directory(
            self.pdf_dir,
            ".pdf",
            self.extract_text_from_pdf_pdfplumber
        )
        all_texts.extend(pdf_texts)
        self.stats['pdf_count'] = len(pdf_texts)

        # Process DOCX files
        docx_texts = self.process_directory(
            self.docx_dir,
            ".docx",
            self.extract_text_from_docx
        )
        all_texts.extend(docx_texts)
        self.stats['docx_count'] = len(docx_texts)

        # Process TXT files
        txt_texts = self.process_directory(
            self.txt_dir,
            ".txt",
            self.extract_text_from_txt
        )
        all_texts.extend(txt_texts)
        self.stats['txt_count'] = len(txt_texts)

        # Add glossary context
        glossary_context = self.add_glossary_context()
        if glossary_context:
            all_texts.insert(0, glossary_context)

        # Combine all texts
        final_text = "\n\n".join(all_texts)

        # Write to output file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_text)

        # Print statistics
        print("\n" + "="*80)
        print("Processing Complete!")
        print("="*80)
        print(f"Total documents processed: {self.stats['total_documents']}")
        print(f"  - PDF files: {self.stats['pdf_count']}")
        print(f"  - DOCX files: {self.stats['docx_count']}")
        print(f"  - TXT files: {self.stats['txt_count']}")
        print(f"Total characters: {self.stats['total_chars']:,}")
        print(f"Output file: {output_path}")
        print(f"File size: {output_path.stat().st_size / 1024:.2f} KB")

        if self.stats['errors']:
            print(f"\nWARNING: Errors encountered: {len(self.stats['errors'])}")
            print("First 5 errors:")
            for error in self.stats['errors'][:5]:
                print(f"  - {error}")

        print("="*80)

        return str(output_path)


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description='Process BEP documents for ML training'
    )
    parser.add_argument(
        '--output',
        type=str,
        default='data/consolidated_training_data.txt',
        help='Output file path for consolidated training data'
    )
    parser.add_argument(
        '--data-dir',
        type=str,
        default='data/training_documents',
        help='Directory containing training documents'
    )

    args = parser.parse_args()

    # Create processor and run
    processor = DocumentProcessor(data_dir=args.data_dir)
    output_file = processor.process_all_documents(output_file=args.output)

    print(f"\nTraining data ready at: {output_file}")
    print("\nNext steps:")
    print("1. Review the consolidated training data")
    print("2. Run training: python scripts/train_model.py --data-file data/consolidated_training_data.txt")
    print("3. Test the model with your BEP generator")


if __name__ == "__main__":
    main()
